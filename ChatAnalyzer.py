try:
    from Crypto.Cipher import AES
    import base64, hashlib
except ImportError:
    print("[INFO] Modul pycryptodome tidak ada, fitur enkripsi nonaktif.")

import re
import tkinter as tk
from tkinter import filedialog, scrolledtext, messagebox, simpledialog, ttk
from collections import defaultdict
import matplotlib.pyplot as plt
import mplcursors
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd
from fpdf import FPDF
import openpyxl
from openpyxl.styles import Font
from textblob import TextBlob
import threading
import queue
import time
from datetime import datetime, timedelta
from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
import numpy as np

class WhatsAppAnalyzer:
    def show_chart(self):
        """Show offensive messages per sender chart"""
        if not self.sender_stats:
            print("No offensive messages to display.")
            return
        senders = list(self.sender_stats.keys())
        counts = list(self.sender_stats.values())
        plt.figure(figsize=(8, 6))
        plt.bar(senders, counts, color='red')
        plt.xticks(rotation=45, ha='right')
        plt.title("Jumlah Pesan Kasar per Pengirim")
        plt.tight_layout()
        plt.show()
    
    def show_senders_chart(self):
        """Show total messages per sender chart"""
        if not self.all_senders:
            print("No messages to display.")
            return
        senders = list(self.all_senders.keys())
        counts = list(self.all_senders.values())
        plt.figure(figsize=(8, 6))
        plt.bar(senders, counts, color='blue')
        plt.xticks(rotation=45, ha='right')
        plt.title("Jumlah Pesan per Pengirim")
        plt.tight_layout()
        plt.show()
    
    def show_metrics_chart(self):
        """Show evaluation metrics chart (precision, recall, f1-score)"""
        metrics = self.calculate_precision_recall()
        
        labels = ['Precision', 'Recall', 'F1-Score']
        values = [metrics['precision'], metrics['recall'], metrics['f1_score']]
        colors = ['#4CAF50', '#2196F3', '#FF9800']
        
        plt.figure(figsize=(8, 6))
        bars = plt.bar(labels, values, color=colors)
        plt.title('Metrik Evaluasi Deteksi Bullying')
        plt.ylabel('Nilai')
        plt.ylim(0, 1)
        
        # Add value labels on bars
        for bar, value in zip(bars, values):
            plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.01, 
                    f'{value:.2f}', ha='center', va='bottom')
        
        plt.tight_layout()
        plt.show()
    
    def show_word_frequency_chart(self):
        """Show word frequency chart for offensive words"""
        if not self.word_stats:
            print("No offensive words to display.")
            return
        
        # Get top 20 offensive words
        sorted_words = sorted(self.word_stats.items(), key=lambda x: x[1], reverse=True)[:20]
        words = [word[0] for word in sorted_words]
        counts = [word[1] for word in sorted_words]
        
        plt.figure(figsize=(10, 8))
        bars = plt.barh(words, counts, color='red')
        plt.title('Frekuensi Kata Kasar')
        plt.xlabel('Jumlah Kemunculan')
        plt.gca().invert_yaxis()  # Display highest count at the top
        
        # Add value labels on bars
        for bar, count in zip(bars, counts):
            plt.text(bar.get_width() + 0.3, bar.get_y() + bar.get_height()/2, 
                    f'{count}', ha='left', va='center')
        
        plt.tight_layout()
        plt.show()
    
    def show_sentiment_chart(self):
        """Show sentiment analysis chart per user"""
        if not self.sentiment_stats:
            print("No sentiment data to display.")
            return
        
        # Prepare data for chart
        senders = []
        positive_scores = []
        neutral_scores = []
        negative_scores = []
        
        for sender, sentiments in self.sentiment_stats.items():
            if sentiments:
                senders.append(sender)
                # Calculate average scores
                avg_positive = sum(s['positive'] for s in sentiments) / len(sentiments)
                avg_neutral = sum(s['neu'] for s in sentiments) / len(sentiments)
                avg_negative = sum(s['negative'] for s in sentiments) / len(sentiments)
                
                positive_scores.append(avg_positive)
                neutral_scores.append(avg_neutral)
                negative_scores.append(avg_negative)
        
        if not senders:
            return
            
        # Create stacked bar chart
        x = np.arange(len(senders))
        width = 0.8
        
        fig, ax = plt.subplots(figsize=(12, 8))
        p1 = ax.bar(x, positive_scores, width, label='Positif', color='#4CAF50')
        p2 = ax.bar(x, neutral_scores, width, bottom=positive_scores, label='Netral', color='#2196F3')
        p3 = ax.bar(x, negative_scores, width, bottom=np.array(positive_scores) + np.array(neutral_scores), 
                   label='Negatif', color='#F44336')
        
        ax.set_ylabel('Proporsi Sentimen')
        ax.set_title('Analisis Sentimen per Pengirim')
        ax.set_xticks(x)
        ax.set_xticklabels(senders, rotation=45, ha='right')
        ax.legend()
        
        # Add value labels
        for i, (pos, neu, neg) in enumerate(zip(positive_scores, neutral_scores, negative_scores)):
            total = pos + neu + neg
            if total > 0:
                ax.text(i, total + 0.01, f'{total:.2f}', ha='center', va='bottom')
        
        plt.tight_layout()
        plt.show()

    def __init__(self):
        # List of offensive words
        self.bad_words = [
            r'\bkontol\b', r'\bk0ntol\b', r'\bk\*ntol\b', r'\bkntl\b', r'\bk0ntl\b', r'\bk\*ntl\b', r'\bKONTOL\b', r'\bK0NTOL\b', r'\bK\*NTOL\b',
            r'\bmemek\b', r'\bm3mek\b', r'\bmmk\b', r'\bm\*m3k\b', r'\bMEMEK\b', r'\bM3MEK\b', r'\bM\*M3K\b',
            r'\banj\b', r'\banjg\b', r'\bajg\b', r'\b4jg\b', r'\ba\*jg\b', r'\banj1ng\b', r'\banj1ngg\b',
            r'\bngentod\b', r'\bng3ntod\b', r'\bngent0d\b', r'\bngentot\b', r'\bng3ntot\b', r'\bng3nt0t\b', r'\bngtd\b', r'\bngt0d\b', r'\bngt\*d\b',
            r'\basu\b', r'\ba\$u\b', r'\ba_su\b', r'\b@su\b',
            r'\bjembut\b', r'\bjmbt\b', r'\bj3mbut\b', r'\bj\*mbut\b', r'\bj3mb\*t\b',
            r'\bpukimak\b', r'\bpuk1mak\b', r'\bpuk1m4k\b', r'\bpkmk\b', r'\bppkmk\b', r'\bp\*kmak\b', r'\bpukim4k\b',
            r'\bnjing\b', r'\bnjingg\b', r'\bnj1ng\b', r'\bn\*j1ng\b',
            r'\bbangsad\b', r'\bbangsat\b', r'\bb@ngsat\b', r'\bb@ngsad\b', r'\bbgst\b', r'\bbgsd\b', r'\bb\*gsd\b', r'\bb@gsd\b',
            r'\banak miskin\b', r'\banak m1skin\b', r'\banak m\*skin\b', r'\bMISKIN\b', r'\bM1SKIN\b', r'\bM\*SKIN\b',
            r'\bidiot\b', r'\b1diot\b', r'\b1d10t\b', r'\b1d1ot\b', r'\bid10t\b', r'\b1d\*ot\b',
            r'\btolol\b', r'\btlol\b', r'\bt0lol\b', r'\bt\*l0l\b', r'\bt\*l\*l\b',
            r'\bjelek\b', r'\bj3lek\b', r'\bjelekk\b', r'\bj3l3k\b',
            r'\bkrempeng\b', r'\bkr3mpeng\b', r'\bkrmpeng\b', r'\bkremp\*ng\b', r'\bkr3mp\*ng\b',
            r'\bpesek\b', r'\bp3sek\b', r'\bp\*sek\b',
            r'\bcacat\b', r'\bc4cat\b', r'\bc\*c4t\b', r'\bc4c4t\b', r'\bcac4t\b',
            r'\blonte\b', r'\bl0nte\b', r'\bl\*nte\b', r'\blont3\b',
            r'\bdasar wanita murahann\b', r'\bdasar wanita murahan\b', r'\bd4s4r wanita murahan\b', r'\bdasar wanita mur4han\b',
            r'\bpelacur\b', r'\bp3lacur\b', r'\bp3l4cur\b', r'\bp\*l4cur\b', r'\bpel4cur\b',
            r'\bcina lu\b', r'\bc1na lu\b', r'\bc1n4 lu\b', r'\bcina l0\b', r'\bcina lo\b',
            r'\bpribumi goblok\b', r'\bpribumi g0blok\b', r'\bpribumi g\*blok\b', r'\bprijbumi gobl0k\b',
            r'\bbencong\b', r'\bb3ncong\b', r'\bb3nc0ng\b', r'\bb\*ncong\b',
            r'\bhama\b', r'\bh4ma\b', r'\bh@ma\b',
            r'\bwaria\b', r'\bwar1a\b', r'\bw4ria\b', r'\bw@ria\b',
            r'\bbanci\b', r'\bb4nci\b', r'\bb4nc1\b', r'\bb@nci\b',
            r'\bmati aja lu\b', r'\bmati aj4 lu\b', r'\bmati a\*\* lu\b', r'\bm4ti aj4 lu\b',
            r'\bgw bunuh lu\b', r'\bgue bunuh lu\b', r'\bgwe bunuh lu\b', r'\bgw b\*nuh lu\b', r'\bgw bunuh l0\b',
            r'\bgw cari rumah lu\b', r'\bgue cari rumah lu\b', r'\bgw c4ri ruma lu\b', r'\bgw c@ri ruma lu\b',
            r'\banak haram\b', r'\banak h4ram\b', r'\banak h@ram\b', r'\banak h\*r4m\b',
            r'\borang tua bodoh\b', r'\borang tua b0doh\b', r'\borang tua b\*doh\b',
            r'\bdasar kafir\b', r'\bd4s4r kafir\b', r'\bdasar k4fir\b', r'\bd@sar k4fir\b',
            r'\bmunafik lu\b', r'\bmun4fik lu\b', r'\bm\*n4fik lu\b', r'\bmuna-fik lu\b',
            r'\bgendut kaya babi\b', r'\bgendut ky babi\b', r'\bg3ndut babi\b', r'\bg3ndut k4y4 babi\b',
            r'\bbabu\b', r'\bb4bu\b', r'\bb@bu\b',
            r'\bngemis\b', r'\bng3mis\b', r'\bnge-mis\b', r'\bng\*m1s\b',
            r'\bbajingan\b', r'\bbjngan\b', r'\bb4jingan\b', r'\bb@jingan\b', r'\bb\*jingan\b',
            r'\bjing\b', r'\bj1ng\b', r'\bj!ng\b', r'\bj\*ng\b',
            r'\btod\b', r'\bt0d\b', r'\bt\*d\b',
            r'\bkntol\b', r'\bk\*ntol\b', r'\bknt0l\b',
            r'\bngewe\b', r'\bng3we\b', r'\bng\*w3\b', r'\bng\*w3h\b',
            r'\bcoli\b', r'\bc0li\b', r'\bc\*li\b', r'\bco1i\b', r'\bc@li\b',
            r'\bcomli\b', r'\bc0mli\b', r'\bc\*mli\b',
            r'\bfuck\b', r'\bf\*ck\b', r'\bfck\b', r'\bfucc\b', r'\bfuxk\b',
            r'\bfuck lah\b', r'\bf\*ck lah\b', r'\bfck lah\b', r'\bfuck\b', r'\bf\*ck\b', r'\bfuxk\b', r'\bfvck\b', r'\bfuqq\b',
            r'\bshit\b', r'\bsh1t\b', r'\bs\*it\b', r'\bsht\b',
            r'\basshole\b', r'\b@sshole\b', r'\ba\*shole\b',
            r'\bbitch\b', r'\bb1tch\b', r'\bb\*tch\b', r'\bb!tch\b', r'\bbi7ch\b',
            r'\bbastard\b', r'\bb@stard\b', r'\bb4stard\b',
            r'\bslut\b', r'\bsl\*t\b', r'\bslvt\b', r'\bs1ut\b',
            r'\bwhore\b', r'\bwh0re\b', r'\bw\*ore\b',
            r'\bcunt\b', r'\bc\*nt\b', r'\bcvnt\b',
            r'\bdick\b', r'\bd1ck\b', r'\bd\*ck\b', r'\bd!ck\b',
            r'\bpussy\b', r'\bp\*ssy\b', r'\bpu55y\b', r'\bp\$\$y\b',
            r'\bpenis\b', r'\bp3nis\b', r'\bp\*nis\b',
            r'\bvagina\b', r'\bv@gina\b', r'\bvag1na\b',
            r'\bdildo\b', r'\bd1ldo\b',
            r'\bjerk\b', r'\bj3rk\b', r'\bj\*rk\b',
            r'\bmoron\b', r'\bm0ron\b', r'\bm\*ron\b',
            r'\bstupid\b', r'\bst00pid\b', r'\bstupd\b', r'\bs\*upid\b',
            r'\bidiot\b', r'\b1diot\b', r'\bi\*iot\b',
            r'\bretard\b', r'\br3tard\b', r'\br\*tard\b',
            r'\bgay\b', r'\bg\*y\b', r'\bg4y\b', r'\bg@ay\b',
            r'\bfaggot\b', r'\bf4ggot\b', r'\bf\*ggot\b',
            r'\bnigger\b', r'\bn\*gger\b', r'\bn1gger\b', r'\bn!gger\b',
            r'\bnigga\b', r'\bn1gga\b', r'\bn\*gga\b',
            r'\bhoe\b', r'\bh0e\b', r'\bh\*e\b',
            r'\bskank\b', r'\bsk4nk\b',
            r'\btrash\b', r'\btr4sh\b',
            r'\bugly\b', r'\bug1y\b', r'\bugl\*y\b',
            r'\bfat\b', r'\bf4t\b', r'\bf@t\b',
            r'\bkill your self\b', r'\bkys\b',
            r'\bcoward\b', r'\bc0ward\b',
            r'\bcrybaby\b', r'\bcryb4by\b',
            r'\battention seeker\b', r'\battention whore\b',
            r'\byou’re nothing\b', r'\bjust die\b',
            r'\byou look disgusting\b', r'\bdisgusting\b',
            r'\byou make me sick\b', r'\bsickening\b',
            r'\bdumb\b', r'\bdvmb\b', r'\bd\*mb\b',
            r'\bbraindead\b', r'\bbrain-dead\b', r'\bbacot lo kontol\b', r'\bbacot lu jawir\b'
        ]
        
        self.sentiment_analyzer = SentimentIntensityAnalyzer()
        # List of bad lines (for testing purposes)
        self.bad_lines = [ 
            "kamu memang bodoh sekali",
            "dasar anak haram",
            "mati saja kau",
            "wajahmu sangat menjijikkan",
            "dasar miskin",
            "kamu tidak berguna",
            "pergi dari sini, sampah",
            "dasar pecundang",
            "tidak ada yang menyukaimu",
            "kamu pantas dihina",
            "bodoh sekali kau",
            "dasar pengecut",
            "kamu tidak akan pernah berhasil",
            "wajahmu membuatku mual",
            "dasar pemalas",
            "kamu tidak berguna sama sekali",
            "dasar pembohong",
            "tidak ada yang percaya padamu",
            "kamu selalu gagal",
            "bodoh bet cok",
            "mukak lu kek memek",
            "plenger lu",
            "mukak lu kek kontol",
            "gmna sih mukak lu kek anjing",
            "tolol bet punya otak",
            "ounya otak gak sih",
            "apa sih jawir",
            "dasar pengkhianat",
            "aquilt tolol",
            "Bego bat aquil",
            "dikatain tolol sama kairi😔😔",
            "tai lu karet",
            "kontol",
            "punya aquil mendelep asal lu tau",
            "Wanjir",
            "lu mendelep sampe kek lucinta Luna ya",
            "gw doain punya lu copot kek gini",
            "copot apanya tuh",
            "Taik terbang anj",
            "parah di katain yatim",
            "santai aj",
            "anj luwh",
            "Eh tai ko gua pramuka si",
            "Goblok",
            "jembut",
            "apa puki",
            "bot anjr no 2,3",
            "apsi bego najis anjr",
            "bego bat",
            "apasi bego lu mau ga digituin?",
            "najis baperan",
            "Wanjir",
            "Marpa pengedar narkoba anjir",
            "homo",
            "aoakoww kontol",
            "goblok",
            "kontol",
            "ngapa salah mulu anj gwe",
            "Lah tolol kelas Laen ga di cek",
            "kontoll",
            "minta goblok",
            "serius anjr gw",
            "ribet banget anj nurul khatimah anak x5",
            "kalo gamuat pake nama panggilan aja cukup anjr",
            "Tai kli",
            "calon murid kontol",
            "Ketiduran anjgggg",
            "kok besok dipake juga anjr celana putih nya",
            "Anjasss ayang nya aqila",
            "GAJELAS BEGOO",
            "iye tai",
            "Males bat tai tarik tambang pagi",
            "minim literasi tai",
            "🧕🏻 : Kaos tunik lengan panjang merah, rok putih sekolah, & jilbab segi 4 merah/putih.",
            "tai lu pan",
            "jelek",
            "Lah kalo rangkuman kelompok sendri mah udah anjir",
            "Lah itu kan MTK anjir",
            "nyukur besok aja anjg",
            "Ah kontol tugas kyk tai dh",
            "Anjir kan milikan Napa disitu jadinya Sinar",
            "Anjas",
            "mutasi kontol bukan do",
            "Beneran anjir",
            "Lah apaan anjir motong pinggir doank",
            "jelekan gw kok",
            "kalo lu jelek gua apaan qil",
            "kim pandangan dr orang lain masi keliatan panjang apalagi guru gg nt ggwp ez nob",
            "bio tugas bejibun tai",
            "WOI GOBLOK",
            "Anjir lu smooting jdi gimana dip",
            "ALAY BANGET KWONTOL WKWKEKKWKWKWKWKW KEK FEMBOY ANJR",
            "Anjg",
            "Jigrak anj",
            "anjayyyyyy",
            "Anjir ada asepnya",
            "Anjir iyak belom ada",
            "eh goblok",
            "abis kontol",
            "BIOLOGI ANJJJ",
            "emang tolol ego",
            "bingung gua ama pola fikir guru anj",
            "Iya anj",
            "kontol kalimantan anj",
            "owh anj gw kira",
            "Anjg beda huruf kontol",
            "Anjg",
            "beda di titik doang kontol",
            "Anjg",
            "80 anj gw wkwk",
            "Iya anjg wkwk",
            "Anjir iya",
            "Tiap web beda anj",
            "pelit bat tai",
            "Anj",
            "tai",
            "tadi tai kucing lupa diangkat ama bapak lu noh",
            "anjai alen",
            "nemu aja anjg",
            "Wanjir bing bikin 50 kalimat",
            "maksudnya gimana dah tai",
            "Emang pekaen ulangan anjir",
            "@62895606013927 lengkapin PKN dan nanti ga boleh ulangan anjir",
            "Hah anjg kata siapa digambar",
            "najis",
            "Wanjing digambar",
            "baper deh ah santai ceees",
            "Mana jj yang kau janjikan itu hah",
            "Lu tai",
            "@628118822409 KO RECIL BS GA LOLOS ANJG",
            "tai",
            "infonya setengah anjay",
            "Itu buku nex Dey estimasi Selasa anjr",
            "Tai Ama delpan",
            "@6285176838927 anjg Ama Nurul",
            "Aset anjir",
            "Mampus dah setiap diskusi nyengir Mulu dah tai",
            "Nyengir2 anjg",
            "ngntod kelompok gua cowo bedua doang tai",
            "apaan anjr gaje banget tiba2 baper 😂😂😂",
            "eh goblok co",
            "Wanjir Depok gerimis",
            "Iris segala janji",
            "Tolol idiot",
            "Kontol",
            "Anjg",
            "anjing",
            "diusir gw puki",
            "Anjai",
            "Lebih banyak LDKO gw anjg",
            "di tanjakan",
            "enak bgt tai itu",
            "anjir",
            "Mampus tolol",
            "kontol",
            "akwkwkw aquil diomelin anjg",
            "kocak anj wkwkwkkw sama kek hakim kejadiannya",
            "Yang 8 goblok banget gitu doang gatau",
            "katanya udh anjirlah",
            "au ah puki",
            "eh anjg bagi fisika dong",
            "anjg",
            "mahal bet anjr skrng seribu satu",
            "anj pembinaan bk",
            "duh rambut2 lagi anj lah",
            "anjg",
            "ga anjr gw dikit banget wkwkkw",
            "lah anjir ada bener salah nya💀",
            "Njing ternyata semuanya jelek",
            "Kirain gw gw yg paling jelek",
            "rabunya gua anjg",
            "anjg Pramuka apaan itu jir",
            "gjls tai",
            "iya anjg",
            "ah tai",
            "kontol kontol",
            "bgst ujan anjg",
            "Mendadak bat tai",
            "gajelas anj pramuka",
            "kacau bet eskul nya hafizh di katain",
            "eh kontol gw cuma beli 4",
            "paling bentuk kontol",
            "memek tu kimmek",
            "oiya lu di palas anjg",
            "anjg",
            "anjg paes",
            "anjg banyak bat",
            "bawa juga anj cadangan",
            "parah bat anjir",
            "kontolllllll",
            "masalahnya gua kejebak anjing quil",
            "lanjutt",
            "kontol",
            "besok ribet bat tai",
            "geo anjg",
            "tolong lah puki",
            "tolol",
            "anjg dah untung Masi rame",
            "anjir berak",
            "mingdepnya ulangan lagi anjir",
            "pkn apaan anj",
            "wanjing presentasi + wwc 2",
            "WANJIR",
            "ANJINGG",
            "sama anjing",
            "aamiin anjg",
            "ngasal ae ngasal anjg",
            "santai mas",
            "ahhh anj ribet banget nih osis",
            "wanjiir gerimis",
            "batre mic 1 pak 25 anjg",
            "tolol",
            "goblok kata gw mah",
            "serius kontol gw gatau apa apa",
            "bikin kontol",
            "Ih ngomong nya kontol kontolan",
            "kontol kontol",
            "anjeng",
            "tai",
            "wibu anj",
            "taii",
            "kontol",
            "yatim",
            "baru 2 bab bloman lanjut",
            "lanjutannya",
            "kacamata idiot",
            "behel idiot",
            "lagian nge fuck fuck diliatin umum tolol",
            "anjg di kotu",
            "gblk anjr di panggil bang",
            "ANJJJJRRR JADI VIEWS TERBANYAK CO",
            "anjir rumah hantu kotu yang gw ga mau masukin",
            "wanjir Nasi goreng kambing etawa bangst <Pesan ini diedit>",
            "wanjing gw lolos lagi",
            "anjir kebanyakan filter",
            "anjg kagak",
            "belom fix ini mah tai",
            "apasi idiot",
            "gua dari dlu pgn banget megang program lomba antar sekolah anjing",
            "wkwkw tolol",
            "gw ditanya mau bikin proker apaan anjg",
            "tp disuruh eval nya dari kaka kakanya tolol bkn program nya aowkwk",
            "fisika gimana dah anj",
            "cewe gw siapa anjf",
            "wanjing",
            "ldks apaan anjg ama guru",
            "wanjir harus potong dulu donk",
            "aowkwowk kontol",
            "serious co wkwkwk diceritain alasan ngapa ga semua murid dpt ldks ama kepala sekolahnya",
            "alwkwowkw kontol",
            "kontol",
            "panjang nya",
            "anjing",
            "apaan anjir dikirik semua",
            "anjayyy",
            "anjr vn desah",
            "siapa ini anjay",
            "yg bener mana anjg",
            "tai itu lu nonton anj",
            "just rp kontol",
            "anjg alajar",
            "skola jelek",
            "wkwkwk anjing emang kan",
            "yaudh datain aja kita ber4",
            "wanjir ngikut² Bu diyah",
            "kontol",
            "mles bet anjg",
            "kalo ga menang remed lagi gitu anjg?",
            "ribet bat dah guru ni anjg",
            "Oh oke dah trobos aja anjing",
            "Santai Napa semua orang juga capek",
            "anjenggggggh",
            "robek anjr",
            "anj pake ada diary lg",
            "Anjir foto ni pala gw kan mobil",
            "nadip tai",
            "Alisha ga lu tai tai in bel?",
            "ga cowo ga di tai2 in",
            "Tai tai in Hakim bel",
            "bego aquil",
            "ad lanjutan nya",
            "rudolf kaga live kontol",
            "bngst anj",
            "anj bugil gw",
            "anjay",
            "kontol",
            "kontol",
            "pernak pernik apaan tai",
            "nitip sang taii",
            "post it apaan anjg",
            "santai",
            "iya apa anjg pernak pernik",
            "gajelas anj bilang nya pernak-pernik",
            "lagian nanya anjg",
            "anjg",
            "kontol",
            "Fisika ulangan di kelas apa di lab anjg",
            "manik manik anjingg",
            "AOWKWOWKWKWKKWKW SANTAI KONTOL",
            "anj",
            "elu goblokk",
            "elu goblok Ama hakim",
            "gw ketawain kegoblokan sarip",
            "pr apaan aja tai",
            "anjg gw kira drtdi yang LKS dikira tugasternyata nanya",
            "bikin gelang anjir",
            "kontol",
            "jadi ngerjain geo bareng ga kwkw anjg",
            "woi gw masi dijalan anj",
            "anjg materinya prinsip juga",
            "meet anj",
            "lah tai",
            "aduh puki males bat ngeprint",
            "septimemek",
            "ribet bat tai",
            "we kontol",
            "ah kontol",
            "ngatur apaan anjir",
            "tapi panjang bt co",
            "anjaoy",
            "nyuruh apa tolol",
            "mau nanya rt jg udh malem anjr",
            "Kan katanya lu emang citayem yaudh cari aja, knp bilang hakim nyuruh² anjir",
            "WOI KONTOL UDAH EGO JANGAN KONFLIK MEMEK",
            "anjg jaman penjajahan",
            "rumah anjir",
            "gada hubungan nya kontol",
            "gw gampang anjir",
            "panjang bet anjg depok tai",
            "yelah kim udah dijelasin panjang lebar ama akhdan",
            "OPSI: anjir (‎5 suara)",
            "OPSI: gatau? anjirr , gatau kenapa bisa (‎5 suara)",
            "OPSI: bisa? anjir bisa tumbuhan? (‎5 suara)",
            "OPSI: 10? banyak anjir (‎5 suara)",
            "kontol",
            "wah anjg lupa ngeprin",
            "udh mulai anjir Brifingnya",
            "iya anjir",
            "goblok iya juga anjg kenapa ga sempet²in ya",
            "anjir iya juga kenapa ga Poto",
            "tai",
            "anjay",
            "badai anjir",
            "siapa anjr yg poto",
            "anjg",
            "kalo jelek nyesel gw nanti",
            "gegara emak² anjing emang",
            "parah bat emak dikatain anjg",
            "kontol",
            "anjg k4 juga bgst",
            "wkwkwk anjg",
            "anj beneran semua",
            "Jadi kita latihan lanjut rapat",
            "goblok",
            "lah kocak anjir",
            "anjg",
            "anjg besok gw Ampe jam 4",
            "ajg sama kayk waktu sekolah tai",
            "PAGI BANGET YAK ANJ",
            "jam 8 kek anj",
            "ebuset jam 6 ngapain anjir",
            "anjg kok dapet",
            "langsung ulangan anj",
            "memek",
            "Gw juga belom belanja part part buat fisika lagi 😂",
            "disuruh lurus anjg",
            "memek",
            "memek",
            "direkam anj kalo lawan arah jiga",
            "memek",
            "tolol",
            "tolol bat anjg",
            "lawan 25 seri bego",
            "anjay",
            "cuma burem bet anj kalo tag² gitu",
            "dah telanjur minta aquil aku mas",
            "kontoll",
            "anjg",
            "Barab lanjutin besok kan",
            "anjg gw juga",
            "kok ada yang ganjel gitu kemaren Bu fad Jumat itu baru rencana",
            "ngantuk kontol",
            "Ihsan mode rambut panjang",
            "@6285176838927 mintain pan ke yg udah",
            "tai apa?, serius ga?, parah anjir kalo iya",
            "wahhh anjirrr",
            "anjir ga expect gw",
            "Eh anjirlah",
            "kontol",
            "laper kontol",
            "ngentod memek",
            "Goblok alquran gw dikelas",
            "manjat pohon aja kaga kynya",
            "NGENTOD NIH 10.1 ANJINGGGGG",
            "wkwk tolol marpa",
            "gw juga mau anj",
            "toy story tai",
            "iya anj itu dia",
            "kontoll",
            "anjir rambut lu Masi keren jay",
            "APAAN SI ANJ ABSURD BET TIBA2 ADA YG SIRIK WKWKKEE 😹",
            "wanjir",
            "anjg gw di bantai",
            "aku nanya doank beb santai donk wkwkwk",
            "santai Di Mpk udh ada yang mau sp 1 OSIS juga bos wkwk",
            "ANJG GUA INI AKWOWKOWKW",
            "tai apa",
            "wkkwkwkwkw donfo hapis tolol",
            "plottwist anjir",
            "list kontol",
            "najis",
            "anj",
            "anjing",
            "anjg tulisannya jam 6 goblok",
            "soalnya jam 6 anj ga impossible co",
            "anj co ga makan 😭",
            "anj",
            "wanjay",
            "dah lama anj ga begini panas bet",
            "anjg",
            "@6281224207592 @6285774628743 parah dikatain titan gagah",
            "lah anj kata lu Titan AOT ANJG",
            "WKWK ANJIR",
            "yg bener tapi anj",
            "anjg di xteler",
            "udah jaga2 anj",
            "apaan anjg Rorrrr",
            "gw gak minta anj",
            "anj",
            "Anjir HAHAHAHAHA",
            "anj beda banget",
            "Sombong memek lu",
            "Najiss",
            "anjayy",
            "Anjay delpan",
            "anjir",
            "gw baru bangun tai",
            "anjg",
            "wah tai",
            "kepenuhan anjir",
            "anj baru bangun hw",
            "takbiratul ihram anj",
            "taik",
            "anjir mar kalo bener ganteng lu",
            "ini lagi tambah goblok",
            "kasti goblok",
            "anjg",
            "Intel bego dia",
            "anjir gw lagi nyari stiker",
            "anj suara siapa",
            "anj",
            "Anjayy",
            "wanjir",
            "Tai",
            "anjay",
            "anjay",
            "goblok",
            "anjg",
            "anjir",
            "anja aokwowwkw",
            "anjir lah",
            "anjr kok keren rel",
            "walah anjing mau ldks malah cek rambut",
            "Pepekkkk",
            "bil lu 2005 anj?",
            "anjir lah",
            "kayak palkon anj",
            "udh pendek anjir ini",
            "tapi iya si kemarin kata ridho ada yg rambut nya panjang malah",
            "biasanya emang ga dicek anjir",
            "anjir gw 3 ini",
            "itu ukuran paling panjang",
            "tai kena ni gua bsj",
            "anjir la",
            "anjir",
            "ANWJSNEUDHEUSHSBS ANJG",
            "iya anjir itu tipis",
            "hakim sendiri Masi panjang",
            "banyakin dah foto panjang nanti botak sebulan",
            "sem 2 awal panjang lagi gw sans aja",
            "iya tapi kayak abis anjiir itu",
            "masi panjang cih",
            "pa nabil tai",
            "aku Masi panjang dikit rip 😔",
            "ga botak yatim",
            "anjg udah",
            "anj gw kira udah botak",
            "owkwowo yatim",
            "ga botak yatim piatu",
            "kaos gambar apa tu anjir",
            "anjir yaqult",
            "anj gw jaga wm botak",
            "Kontol",
            "kontol0",
            "itu tugas maksudnya gimana anjer",
            "apaan anj",
            "gw buat drama kocak anjir",
            "aduh anjir lah yang nilai optimum yak",
            "apala puki nilai optimum",
            "si goblok",
            "tai dah",
            "seluruh kelas anjer",
            "kelebihan anj",
            "anjir",
            "tpi gw goblok",
            "taik",
            "gua baru kelar mandi anjing",
            "mana banjir lagi warsil",
            "anjg la",
            "gua tungguin gamuncul muncul anj",
            "besok pulang abis zubur ato lanjut kbm dah?",
            "tai ah",
            "anjir",
            "OH ANJ gw liat di getcontact trnyt lu yg sv no gw pake kek gitu",
            "dah di sekolah anj",
            "bawa ga anj",
            "Anjay Sarip kembar seiras beda usia",
            "najis mending yg laen",
            "glossy memek",
            "Anjir lahh",
            "tolol",
            "jorok anj ni orang",
            "kontolll",
            "mahal anj api ga ad yg mati",
            "anj Pas Cok",
            "gosah dibahas memek",
            "dari ngemper dibawah tiba tiba diatas anjir ngobrol mulu",
            "ah anj lah hati2 dah kalo ama ortu friendly",
            "dah bngst kagak usah dibahas anj malu gw",
            "mana ada anjinf",
            "pas awal anj",
            "iya anj",
            "Ah memek dahh",
            "anjir masuk cok",
            "Gua kga boleh anj ah",
            "boleh ngga asal yg lain pd gamasuk jg anj",
            "sama anj",
            "gw masuk tuh ada alesan laen anjoy😊",
            "gua juga anjg",
            "tau anjr ga masuk sekola mala di ambil hp gw",
            "anjir wkwk",
            "malah turu si anj",
            "anjg",
            "najis ah ajg",
            "makan makan mulu kontol",
            "berat anjg",
            "alay bat dah lu tai",
            "deket tai",
            "anjir",
            "JEMBUTTTT GW COWO SENDIRI KELAS 10",
            "anjg",
            "anjir sang",
            "sama aja puki",
            "anjir urgent bat",
            "jangan anj ya sang🙏🏻🙏🏻",
            "anj banget gw dapet paling depan",
            "anjg enak bet",
            "lantai 2 yang udah keluar dong",
            "tai apa",
            "anj ngapain elu",
            "siapa anjg?",
            "anjg",
            "mending pukis",
            "Anggap saja fisika itu seperti pria santai",
            "anjg iya sigma co",
            "fisika dan bahasa arab datang tapi kamu ingat kalo kamu adalah pria santai",
            "anj",
            "dikirim anjg",
            "lu gede anjir rip",
            "Kontoll",
            "Gua the real ngasal memek",
            "anj gw 77",
            "memek",
            "anjing",
            "ini lwbih anjg ege",
            "Kok 48 anjai",
            "jelek2 bangsat nilai nya wkwkwkw",
            "iya taik",
            "tai emang",
            "ANJING AOWKWOW KIMIA",
            "YG PENTING GA NYONTEK AH ANJ",
            "lagian pada tolol banget anjg",
            "lagian mau nyontek juga gmn anj",
            "Gua ngerjain nya di kelas ya tai",
            "kalo ganjil naikin, kalo genap gausa",
            "yah mintain lagi kim",
            "anjg💀💀💀",
            "bangsat gw awalnya liat lu 85 tai",
            "aowkw kontol",
            "gjls ajg sejarah pukimai",
            "kok bisa si anjir kyk orang pinter jelek",
            "Mati lah gua anjay",
            "mtk Ama Sosio anj",
            "masa jelek",
            "sok sok pinter bat dikelas tai blg gampang gampang",
            "cara ngitungnya gmn tai",
            "Ah pepek",
            "bilang be jelek",
            "barang ilang beda anj",
            "ohiya anj",
            "nafisatul jahim anj",
            "tai nih reja arap",
            "hakim anjg",
            "apaansi kontol kok fisik",
            "reja arap kan jelek",
            "goblok",
            "owh anjer lah",
            "kata gua anjg",
            "kontol marva",
            "anjg",
            "dahlah remed semua yatim",
            "Lu gimana sih waktu itu lu kan yg ngatain b arab gw jelek",
            "Itu bahasa arab apa anjg",
            "anjg",
            "gausa ngungkit tai",
            "sama aja idiot",
            "debat apaansi kontol",
            "disemangatin malah kek gitu anjing",
            "idiot",
            "debat apaan idiot",
            "lu nya aja nanggepin nya keras goblok",
            "kontol",
            "anjir",
            "kepencet anjg",
            "Irgi goblok",
            "kontol",
            "kontol",
            "tai lu ah",
            "kontol",
            "bangsat udah kelar itu anj",
            "anj gw sekarang gabisa co",
            "wanjir",
            "kontollllll",
            "beda aja semua anjg di buku 93",
            "goblok ajg",
            "anj basket",
            "kontol osis",
            "tai",
            "anjg perkap ikutan lomba",
            "kontol gjls",
            "ms panit gabole main si kontol ah",
            "itu napa ada logo bangor anjg",
            "anj itu yang di buku soal kemaren?",
            "anjr",
            "boong anj ternyata",
            "wanjir",
            "anjay wkkw",
            "ez aj ekontol",
            "ah tai",
            "anjay",
            "tolol",
            "Anjirrr",
            "demi apa msi lanjut kbm",
            "goblok",
            "memek",
            "kimia yg sebelumnya dah abis anjir masa claimny",
            "Anjayy",
            "wajib tolol",
            "tai kucing",
            "gaada baju panjang co",
            "duh malu gw anjim",
            "iya anjg",
            "autis bego tbtb ngirim sound",
            "maksud gua nyantai dulu",
            "tai",
            "anjayyyy",
            "typo anjg",
            "lu Uda jauh ngapain anjir",
            "botak panjang",
            "Bunga makam anj ler",
            "minta anj",
            "bego",
            "tai lah co",
            "lebih tai tes minat bakat",
            "WOI TAI",
            "tai da",
            "ga gitu anj",
            "kontol",
            "anjer",
            "kontol",
            "anjay pada bangun pagi eni",
            "gatau anjir",
            "anj",
            "mana anjoy",
            "dimana anjg",
            "anjir kuas malem²",
            "banyak anjir",
            "tai",
            "bukan kurang anjg",
            "tutup anj",
            "jet siapa anj",
            "bayarnya besom ya anjg",
            "jijik anjg",
            "anjay",
            "goblok",
            "anjir",
            "santai sang",
            "cewe mah gada, someone to talk nya masi lanjut",
            "anjir tau aja lu",
            "kontol",
            "kontol anak ajg",
            "geografi anjing",
            "anj",
            "sejarah gua bloman anjas",
            "anjir kelas sang",
            "belom anj",
            "yang bener yang mana anj",
            "sholat tai main hp mulu",
            "anjir besok",
            "anj papoy aowkwowkowkw",
            "darimana biologi nya pepek",
            "enak bat udah anjay",
            "mati bego mati",
            "ngundang kemokelan anjg ini",
            "mending kga usah masuk tai",
            "kamis jumat nanggung anjing",
            "KONTOL",
            "anjir iya yak",
            "gajelas anj",
            "anjg",
            "pukimay",
            "Kalo gw strategi nya tulisan gw,gw jelekin sejelek jeleknya",
            "gaperlu dijelekin pan💀",
            "nama Alex anj jelek bet",
            "kontol",
            "santai",
            "Anjay",
            "lanjut in",
            "nama lu dah ditulis bego",
            "Gw udh janji dari sebelum Ramadan jir",
            "lu gua tanya berapa ga dijawab anjg",
            "ghost touch ke tai dah",
            "anjay Irgi",
            "gaada anjir",
            "eh ajg yg gamisan siapa aja tai",
            "dc dc kontol",
            "santai",
            "TAI BET ANJGG",
            "iya anjg"
            ]

         # Gabungkan dengan bad_words untuk meningkatkan deteksi
        self.bad_words.extend([re.escape(line) for line in self.bad_lines])
        self.pattern = re.compile('|'.join(self.bad_words), flags=re.IGNORECASE)

        # Combine all patterns with OR
        self.pattern = re.compile('|'.join(self.bad_words), flags=re.IGNORECASE)
        # Pattern for WhatsApp timestamp format
        self.chat_pattern = re.compile(
            r'^(\[?)(\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4},? \d{1,2}[:.]\d{2}(?::\d{2})?(?: [AP]M)?\]?) - (.*?): (.*)$'
        )
        self.offensive_messages = []
        self.sender_stats = defaultdict(int)  # Count of offensive messages per sender
        self.all_senders = defaultdict(int)  # Total count of messages per sender
        self.word_stats = defaultdict(int)
        self.message_timestamps = defaultdict(list)  # Timestamps for each sender
        self.has_data = False
        self.anonymized = False
        self.anonymized_names = {}
        self.sentiment_stats = defaultdict(list)  # Store sentiment scores per sender

    def parse_chat(self, filename, start_date=None, end_date=None):
        line_count = 0
        match_count = 0
        with open(filename, 'r', encoding='utf-8', errors='replace') as f:
            print(f"File berhasil dibuka: {filename}")
            for line in f:
                line_count += 1
                line = line.strip()
                if not line:
                    continue
                print(f"Baris #{line_count}: {line}")
                match = self.chat_pattern.match(line)
                if match:
                    match_count += 1
                    try:
                        timestamp = match.group(2)
                        sender = match.group(3)
                        message = match.group(4)
                        print(f" > Format dikenali: [{timestamp}] {sender}: {message}")
                        # Convert timestamp to datetime
                        dt = self._parse_timestamp(timestamp)
                        if start_date and dt < start_date:
                            continue
                        if end_date and dt > end_date:
                            continue
                        # Add sender to all senders list
                        self.all_senders[sender] += 1
                        # Check message
                        self._check_message(message, sender, timestamp)
                        # Perform sentiment analysis
                        self._analyze_sentiment(message, sender)
                    except IndexError:
                        print(" > Error parsing line:", line)
                else:
                    print(f" > Format tidak dikenali: {line}")
        print(f"Total baris: {line_count}, Format dikenali: {match_count}")
        # Set flag if data is detected
        if len(self.offensive_messages) > 0:
            self.has_data = True

    def _parse_timestamp(self, timestamp):
        # Try different formats
        formats = [
            '%d/%m/%Y, %H:%M',       # dd/mm/yyyy, HH:MM
            '%d-%m-%Y, %H:%M',       # dd-mm-yyyy, HH:MM
            '%d.%m.%Y, %H:%M',       # dd.mm.yyyy, HH:MM
            '%d/%m/%y, %H:%M',       # dd/mm/yy, HH:MM
            '%d-%m-%y, %H:%M',       # dd-mm-yy, HH:MM
            '%d.%m.%y, %H:%M',       # dd.mm.yy, HH:MM
            '%d/%m/%Y %H.%M',        # dd/mm/yyyy HH.MM
            '%d-%m-%Y %H.%M',        # dd-mm-yyyy HH.MM
            '%d.%m.%Y %H.%M',        # dd.mm.yyyy HH.MM
            '%d/%m/%y %H.%M',        # dd/mm/yy HH.MM
            '%d-%m-%y %H.%M',        # dd-mm-yyyy HH.MM
            '%d.%m.%y %H.%M',        # dd.mm.yy HH.MM
        ]
        for fmt in formats:
            try:
                return datetime.strptime(timestamp, fmt)
            except ValueError:
                continue
        raise ValueError(f"Timestamp format tidak dikenali: {timestamp}")

    def _check_message(self, message, sender, timestamp):
        print(f"Memeriksa pesan: {message}")
        try:
            # Simple method - check one by one
            matches = []
            for word in self.bad_words:
                pattern = re.compile(word, re.IGNORECASE)
                if pattern.search(message):
                    # Extract the exact word found
                    found_word = pattern.search(message).group(0)
                    matches.append(found_word)
            # Debug info
            if matches:
                print(f"TERDETEKSI! Kata kasar dalam '{message}': {matches}")
                # Update statistics
                for word in matches:
                    self.word_stats[word.lower()] += 1
                self.sender_stats[sender] += 1
                # Save problematic message
                self.offensive_messages.append({
                    'sender': sender,
                    'time': timestamp,
                    'message': message,
                    'bad_words': matches
                })
            else:
                print(f"Tidak ada kata kasar dalam '{message}'")
        except Exception as e:
            print(f"Error dalam _check_message: {e}")

    def _analyze_sentiment(self, message, sender):
        try:
            # Analisis sentimen dengan VADER
            scores = self.sentiment_analyzer.polarity_scores(message)
            
            # Simpan data sentimen
            self.sentiment_stats[sender].append({
                'compound': scores['compound'],
                'negative': scores['neg'],
                'neu': scores['neu'],
                'positive': scores['pos'],
                'message': message
            })
            
            # Tambahkan skor negatif ke statistik jika terdeteksi sebagai bullying
            if scores['neg'] > 0.5:
                self.sender_stats[sender] += 0.5  # Beri bobot setengah
        except Exception as e:
            print(f"Error in sentiment analysis: {e}")

    def calculate_precision_recall(self):
        """Hitung precision dan recall untuk deteksi bullying"""
        true_positives = 0
        false_positives = 0
        false_negatives = 0
        
        # Hitung berdasarkan bad_lines sebagai ground truth
        for msg_data in self.sentiment_stats.values():
            for sentiment in msg_data:
                message = sentiment['message']
                
                # Cek apakah pesan mengandung kalimat bullying
                is_bullying = any(bad_line in message.lower() for bad_line in self.bad_lines)
                
                # Cek apakah terdeteksi sebagai ofensif
                is_detected = any(msg['message'] == message for msg in self.offensive_messages)
                
                if is_bullying and is_detected:
                    true_positives += 1
                elif is_bullying and not is_detected:
                    false_negatives += 1
                elif not is_bullying and is_detected:
                    false_positives += 1
        
        # Hitung metrik
        precision = true_positives / (true_positives + false_positives) if (true_positives + false_positives) > 0 else 0
        recall = true_positives / (true_positives + false_negatives) if (true_positives + false_negatives) > 0 else 0
        f1_score = 2 * (precision * recall) / (precision + recall) if (precision + recall) > 0 else 0
        
        return {
            'precision': precision,
            'recall': recall,
            'f1_score': f1_score,
            'true_positives': true_positives,
            'false_positives': false_positives,
            'false_negatives': false_negatives
        }


    def anonymize_senders(self):
        """Replace sender names with anonymized versions"""
        if self.anonymized:
            return
            
        self.anonymized_names = {}
        new_all_senders = defaultdict(int)
        new_sender_stats = defaultdict(int)
        new_sentiment_stats = defaultdict(list)
        
        # Create mapping of original names to anonymized names
        for idx, sender in enumerate(self.all_senders.keys()):
            anonymized_name = f"User{idx+1}"
            self.anonymized_names[sender] = anonymized_name
            
            # Update sender counts
            new_all_senders[anonymized_name] = self.all_senders[sender]
            new_sender_stats[anonymized_name] = self.sender_stats.get(sender, 0)
            
            # Update sentiment data
            if sender in self.sentiment_stats:
                new_sentiment_stats[anonymized_name] = self.sentiment_stats[sender]
        
        # Update offensive messages
        for msg in self.offensive_messages:
            if msg['sender'] in self.anonymized_names:
                msg['sender'] = self.anonymized_names[msg['sender']]
        
        # Replace data structures
        self.all_senders = new_all_senders
        self.sender_stats = new_sender_stats
        self.sentiment_stats = new_sentiment_stats
        self.anonymized = True
        
        return self.anonymized_names

    def unanonymize_senders(self):
        """Restore original sender names"""
        if not self.anonymized or not self.anonymized_names:
            return
            
        # Reverse mapping: anonymized name -> original name
        reverse_mapping = {v: k for k, v in self.anonymized_names.items()}
        
        new_all_senders = defaultdict(int)
        new_sender_stats = defaultdict(int)
        new_sentiment_stats = defaultdict(list)
        
        # Restore original names
        for anon_name, count in self.all_senders.items():
            original_name = reverse_mapping.get(anon_name, anon_name)
            new_all_senders[original_name] = count
            new_sender_stats[original_name] = self.sender_stats.get(anon_name, 0)
            
            if anon_name in self.sentiment_stats:
                new_sentiment_stats[original_name] = self.sentiment_stats[anon_name]
        
        # Update offensive messages
        for msg in self.offensive_messages:
            if msg['sender'] in reverse_mapping:
                msg['sender'] = reverse_mapping[msg['sender']]
        
        # Replace data structures
        self.all_senders = new_all_senders
        self.sender_stats = new_sender_stats
        self.sentiment_stats = new_sentiment_stats
        self.anonymized = False
        self.anonymized_names = {}

    def generate_report(self):
        """Create report with data validation"""
        report = []
        total_offensive = len(self.offensive_messages)
        total_words = sum(self.word_stats.values())
        total_senders = len(self.all_senders)
        total_messages = sum(self.all_senders.values())
        report.append("=== LAPORAN ANALISIS PESAN KASAR ===")
        report.append(f"Total pesan bermasalah: {total_offensive} dari {total_messages} pesan")
        report.append(f"Total kata kasar terdeteksi: {total_words}")
        report.append(f"Total peserta dalam chat: {total_senders}")

        # Add sentiment summary
        report.append("\n=== ANALISIS SENTIMEN ===")
        overall_negative = 0
        overall_compound = 0
        sentiment_count = 0

        for sender, sentiments in self.sentiment_stats.items():
            if sentiments:
                avg_negative = sum(s['negative'] for s in sentiments) / len(sentiments)
                avg_compound = sum(s['compound'] for s in sentiments) / len(sentiments)
                overall_negative += avg_negative
                overall_compound += avg_compound
                sentiment_count += 1

                if avg_compound >= 0.05:
                    sentiment_label = "Positif"
                elif avg_compound <= -0.05:
                    sentiment_label = "Negatif"
                else:
                    sentiment_label = "Netral"

                report.append(f"\n{sender}:")
                report.append(f"- Sentimen rata-rata: {sentiment_label} ({avg_compound:.2f})")
                report.append(f"- Skor negatif: {avg_negative:.2f}")

        if sentiment_count > 0:
            overall_negative /= sentiment_count
            overall_compound /= sentiment_count

            if overall_compound >= 0.05:
                overall_label = "Positif"
            elif overall_compound <= -0.05:
                overall_label = "Negatif"
            else:
                overall_label = "Netral"

            report.append(f"\nSentimen keseluruhan: {overall_label} ({overall_compound:.2f})")
            report.append(f"Skor negatif keseluruhan: {overall_negative:.2f}")
        else:
            report.append("\nTidak ada data sentimen yang tersedia")

        # Tambahkan metrik evaluasi
        metrics = self.calculate_precision_recall()
        report.append("\n=== EVALUASI DETEKSI BULLYING ===")
        report.append(f"Precision: {metrics['precision']:.2f} (Akurasi deteksi positif)")
        report.append(f"Recall: {metrics['recall']:.2f} (Kemampuan menemukan semua kasus)")
        report.append(f"F1 Score: {metrics['f1_score']:.2f} (Keseimbangan precision-recall)")
        report.append(f"True Positives: {metrics['true_positives']} (Bullying terdeteksi dengan benar)")
        report.append(f"False Positives: {metrics['false_positives']} (Bukan bullying tapi terdeteksi)")
        report.append(f"False Negatives: {metrics['false_negatives']} (Bullying tapi tidak terdeteksi)")

        # Show all senders and their statistics
        report.append("\n=== STATISTIK PENGGUNA ===")
        sorted_senders = sorted(self.all_senders.items(), key=lambda x: (-self.sender_stats.get(x[0], 0), x[0]))
        for sender, total_msgs in sorted_senders:
            offensive_count = self.sender_stats.get(sender, 0)
            clean_msgs = total_msgs - offensive_count
            offensive_percentage = (offensive_count / total_msgs * 100) if total_msgs > 0 else 0
            report.append(f"\n{sender}:")
            report.append(f"- Total pesan: {total_msgs}")
            report.append(f"- Pesan kasar: {offensive_count} ({offensive_percentage:.1f}%)")
            if offensive_count > 0:
                user_bad_words = defaultdict(int)
                for msg in self.offensive_messages:
                    if msg['sender'] == sender:
                        for word in msg['bad_words']:
                            user_bad_words[word.lower()] += 1
                if user_bad_words:
                    sorted_words = sorted(user_bad_words.items(), key=lambda x: x[1], reverse=True)
                    report.append("Kata kasar yang digunakan:")
                    for word, count in sorted_words:
                        report.append(f" - {word}: {count}x")
                user_offensive_msgs = [msg for msg in self.offensive_messages if msg['sender'] == sender]
                if user_offensive_msgs:
                    report.append("Semua Pesan Bermasalah:")
                    for idx, msg in enumerate(user_offensive_msgs, 1):
                        bad_words = ', '.join(msg['bad_words'])
                        report.append(f"{idx}. [{msg['time']}] {msg['sender']}: {msg['message']}")
                        report.append(f" → Kata terdeteksi: {bad_words}")
                else:
                    report.append("Pengguna ini tidak mengirim pesan kasar.")
            else:
                report.append("Pengguna ini tidak mengirim pesan kasar.")
        if self.word_stats:
            report.append("\n=== STATISTIK KATA KASAR ===")
            report.append("Kata, Frekuensi")
            for word, count in sorted(self.word_stats.items(), key=lambda x: x[1], reverse=True):
                report.append(f"{word}, {count}")
        return '\n'.join(report)
    

class AnalyzerGUI(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("WhatsApp Chat Analyzer v3.0")
        self.geometry("1000x800")
        self.analyzer = WhatsAppAnalyzer()
        self.original_analyzer = None  # To store original data before anonymization
        self.current_filename = None
        self._create_widgets()
        self.monitor_queue = queue.Queue()
        self.monitor_thread = threading.Thread(target=self._monitor_live_chat, daemon=True)
        self.monitor_thread.start()
        
    def _create_widgets(self):
        main_frame = tk.Frame(self)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Control Panel
        control_frame = tk.Frame(main_frame)
        control_frame.pack(fill=tk.X, pady=10)
        
        # Time Period Frame
        time_frame = tk.Frame(control_frame)
        time_frame.pack(side=tk.LEFT, padx=5)
        
        tk.Label(time_frame, text="Periode Waktu:").pack(side=tk.LEFT)
        self.time_period = ttk.Combobox(time_frame, 
                                      values=["Semua Waktu", "1 Bulan Terakhir", "3 Bulan Terakhir", 
                                              "6 Bulan Terakhir", "1 Tahun Terakhir"],
                                      state="readonly",
                                      width=15)
        self.time_period.set("Semua Waktu")
        self.time_period.pack(side=tk.LEFT, padx=5)
        
        # Anonymization Frame
        anon_frame = tk.Frame(control_frame)
        anon_frame.pack(side=tk.LEFT, padx=10)
        
        self.anon_btn = tk.Button(anon_frame, 
                                text="Anonimkan", 
                                command=self._anonymize_data,
                                bg='#9C27B0',
                                fg='white')
        self.anon_btn.pack(side=tk.LEFT, padx=2)
        
        self.unanon_btn = tk.Button(anon_frame, 
                                  text="Batal Anonim", 
                                  command=self._unanonymize_data,
                                  bg='#7B1FA2',
                                  fg='white',
                                  state=tk.DISABLED)
        self.unanon_btn.pack(side=tk.LEFT, padx=2)
        
        # Search Frame
        search_frame = tk.Frame(control_frame)
        search_frame.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        tk.Label(search_frame, text="Cari Pengguna:").pack(side=tk.LEFT)
        self.search_entry = tk.Entry(search_frame, width=30)
        self.search_entry.pack(side=tk.LEFT, padx=5)
        self.search_entry.bind("<Return>", lambda e: self._search_user())
        search_btn = tk.Button(search_frame, 
                             text="Cari", 
                             command=self._search_user,
                             bg='#FF9800',
                             fg='white')
        search_btn.pack(side=tk.LEFT, padx=5)
        
        # Action Buttons
        action_frame = tk.Frame(control_frame)
        action_frame.pack(side=tk.RIGHT, padx=5)
        
        open_btn = tk.Button(action_frame, 
                           text="Buka File Chat", 
                           command=self._open_file,
                           bg='#4CAF50',
                           fg='white')
        open_btn.pack(side=tk.LEFT, padx=5)
        
        chart_btn = tk.Button(action_frame,
                            text="Diagram Pengirim",
                            command=self._show_chart,
                            bg='#2196F3',
                            fg='white')
        chart_btn.pack(side=tk.LEFT, padx=5)
        
        # Chart Buttons Frame
        chart_buttons_frame = tk.Frame(main_frame)
        chart_buttons_frame.pack(fill=tk.X, pady=10)
        
        metrics_btn = tk.Button(chart_buttons_frame,
                              text="Diagram Metrik Evaluasi",
                              command=self._show_metrics_chart,
                              bg='#9C27B0',
                              fg='white')
        metrics_btn.pack(side=tk.LEFT, padx=5)
        
        words_btn = tk.Button(chart_buttons_frame,
                            text="Diagram Frekuensi Kata",
                            command=self._show_word_frequency_chart,
                            bg='#F44336',
                            fg='white')
        words_btn.pack(side=tk.LEFT, padx=5)
        
        sentiment_btn = tk.Button(chart_buttons_frame,
                                text="Diagram Sentimen",
                                command=self._show_sentiment_chart,
                                bg='#FF9800',
                                fg='white')
        sentiment_btn.pack(side=tk.LEFT, padx=5)
        
        # Save Buttons
        save_frame = tk.Frame(main_frame)
        save_frame.pack(fill=tk.X, pady=10)
        
        save_txt_btn = tk.Button(save_frame,
                               text="Simpan ke TXT",
                               command=self._save_to_txt,
                               bg='#607D8B',
                               fg='white')
        save_txt_btn.pack(side=tk.LEFT, padx=5)
        
        save_docx_btn = tk.Button(save_frame,
                                text="Simpan ke DOCX",
                                command=self._save_to_docx,
                                bg='#3F51B5',
                                fg='white')
        save_docx_btn.pack(side=tk.LEFT, padx=5)
        
        save_excel_btn = tk.Button(save_frame,
                                   text="Simpan ke Excel",
                                   command=self._save_to_excel,
                                   bg='#F44336',
                                   fg='white')
        save_excel_btn.pack(side=tk.LEFT, padx=5)
        
        save_pdf_btn = tk.Button(save_frame,
                                 text="Simpan ke PDF",
                                 command=self._save_to_pdf,
                                 bg='#E91E63',
                                 fg='white')
        save_pdf_btn.pack(side=tk.LEFT, padx=5)
        
        save_csv_btn = tk.Button(save_frame,
                                 text="Simpan ke CSV",
                                 command=self._save_to_csv,
                                 bg='#FFC107',
                                 fg='black')
        save_csv_btn.pack(side=tk.LEFT, padx=5)
        
        # Report Area
        self.report_area = scrolledtext.ScrolledText(main_frame,
                                                   wrap=tk.WORD,
                                                   font=('Arial', 10),
                                                   padx=10,
                                                   pady=10)
        self.report_area.pack(fill=tk.BOTH, expand=True)
        self.report_area.config(state='disabled')

    def _search_user(self):
        """Cari pengguna dan tampilkan statistiknya"""
        search_term = self.search_entry.get().strip()
        if not search_term:
            messagebox.showwarning("Peringatan", "Masukkan nama pengguna untuk dicari!")
            return

        matching_senders = [sender for sender in self.analyzer.all_senders.keys()
                            if search_term.lower() in sender.lower()]

        if not matching_senders:
            messagebox.showinfo("Informasi", f"Tidak ditemukan pengguna dengan nama '{search_term}'")
            return

        report = []
        for sender in matching_senders:
            total_msgs = self.analyzer.all_senders.get(sender, 0)
            offensive_count = self.analyzer.sender_stats.get(sender, 0)
            offensive_percentage = (offensive_count / total_msgs * 100) if total_msgs > 0 else 0
            report.append(f"{sender}: Total pesan={total_msgs}, Pesan kasar={offensive_count} ({offensive_percentage:.1f}%)")

        self._update_report_area("\n".join(report))
    
    def _save_to_csv(self):
        """Save report to CSV file"""
        if not self.analyzer.has_data and len(self.analyzer.all_senders) == 0:
            messagebox.showwarning("Peringatan", "Tidak ada data untuk disimpan!")
            return
        file_path = filedialog.asksaveasfilename(defaultextension=".csv",
                                               filetypes=[("CSV Files", "*.csv"), ("All Files", "*.*")],
                                               title="Simpan Laporan sebagai CSV")
        if file_path:
            try:
                import csv
                with open(file_path, 'w', newline='', encoding='utf-8') as f:
                    writer = csv.writer(f)
                    writer.writerow(["Pengirim", "Total Pesan", "Pesan Kasar", "Persentase Kasar"])
                    for sender, total in self.analyzer.all_senders.items():
                        offensive = self.analyzer.sender_stats.get(sender, 0)
                        percentage = (offensive / total * 100) if total > 0 else 0
                        writer.writerow([sender, total, offensive, f"{percentage:.1f}%"])
                messagebox.showinfo("Sukses", f"Laporan berhasil disimpan ke: {file_path}")
            except Exception as e:
                messagebox.showerror("Error", f"Gagal menyimpan file CSV:\n{str(e)}")
    
    def _anonymize_data(self):
        """Anonymize sender names in the data"""
        if not self.analyzer.has_data and len(self.analyzer.all_senders) == 0:
            messagebox.showwarning("Peringatan", "Tidak ada data untuk dianonimkan!")
            return
            
        # Backup original data if not already anonymized
        if not self.analyzer.anonymized:
            import copy
            self.original_analyzer = copy.deepcopy(self.analyzer)
            
        # Perform anonymization
        self.analyzer.anonymize_senders()
        self.unanon_btn.config(state=tk.NORMAL)
        self.anon_btn.config(state=tk.DISABLED)
        
        # Update report
        report = self.analyzer.generate_report()
        self._update_report_area(report)
        messagebox.showinfo("Sukses", "Data berhasil dianonimkan!")

    def _unanonymize_data(self):
        """Restore original sender names"""
        if self.original_analyzer:
            self.analyzer = self.original_analyzer
            self.analyzer.unanonymize_senders()
            self.unanon_btn.config(state=tk.DISABLED)
            self.anon_btn.config(state=tk.NORMAL)
            
            # Update report
            report = self.analyzer.generate_report()
            self._update_report_area(report)
            messagebox.showinfo("Sukses", "Anonimisasi dibatalkan!")

    def _open_file(self):
        filetypes = [('File Chat WhatsApp', '*.txt'), ('Semua File', '*.*')]
        filename = filedialog.askopenfilename(title='Pilih File Chat', filetypes=filetypes)
        if filename:
            self.current_filename = filename
            # Reset previous data
            self.analyzer = WhatsAppAnalyzer()
            
            # Get selected time period
            period = self.time_period.get()
            start_date = None
            end_date = datetime.now()
            
            if period == "1 Bulan Terakhir":
                start_date = end_date - timedelta(days=30)
            elif period == "3 Bulan Terakhir":
                start_date = end_date - timedelta(days=90)
            elif period == "6 Bulan Terakhir":
                start_date = end_date - timedelta(days=180)
            elif period == "1 Tahun Terakhir":
                start_date = end_date - timedelta(days=365)
            
            try:
                self.analyzer.parse_chat(filename, start_date, end_date)
                report = self.analyzer.generate_report()
                self._update_report_area(report)
                
                # Reset anonymization buttons
                self.unanon_btn.config(state=tk.DISABLED)
                self.anon_btn.config(state=tk.NORMAL)
            except Exception as e:
                messagebox.showerror("Error", f"Gagal memproses file:\n{str(e)}")

    def _update_report_area(self, report):
        """Update the report area with new content"""
        self.report_area.config(state='normal')
        self.report_area.delete(1.0, tk.END)
        self.report_area.insert(tk.END, report)
        self.report_area.config(state='disabled')

    def _show_chart(self):
        """Wrapper method to show chart with error handling"""
        if not hasattr(self.analyzer, 'all_senders') or len(self.analyzer.all_senders) == 0:
            messagebox.showinfo("Informasi", "Tidak ada data untuk ditampilkan! Silakan buka file chat terlebih dahulu.")
            return
        # Even if there are no offensive words, we can still show sender chart
        if self.analyzer.has_data:
            self.analyzer.show_chart()
        else:
            self.analyzer.show_senders_chart()  # Show only sender chart

    def _save_to_txt(self):
        """Save report to TXT file"""
        if not self.analyzer.has_data and len(self.analyzer.all_senders) == 0:
            messagebox.showwarning("Peringatan", "Tidak ada data untuk disimpan!")
            return
        file_path = filedialog.asksaveasfilename(defaultextension=".txt",
                                               filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")],
                                               title="Simpan Laporan sebagai TXT")
        if file_path:
            try:
                report_text = self.report_area.get("1.0", tk.END)
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(report_text)
                messagebox.showinfo("Sukses", f"Laporan berhasil disimpan ke:\n{file_path}")
            except Exception as e:
                messagebox.showerror("Error", f"Gagal menyimpan file TXT:\n{str(e)}")

    def _save_to_docx(self):
        """Save report to DOCX (Word) file"""
        if not self.analyzer.has_data and len(self.analyzer.all_senders) == 0:
            messagebox.showwarning("Peringatan", "Tidak ada data untuk disimpan!")
            return
        file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                               filetypes=[("Word Files", "*.docx"), ("All Files", "*.*")],
                                               title="Simpan Laporan sebagai DOCX")
        if file_path:
            try:
                doc = Document()
               # Add title
                title = doc.add_paragraph("=== LAPORAN ANALISIS PESAN KASAR WHATSAPP ===")
                title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                title.runs[0].font.size = Pt(16)
                title.runs[0].font.bold = True
                doc.add_paragraph("")
                # Get text from report area
                report_text = self.report_area.get("1.0", tk.END)
                # Split into paragraphs and add to document
                for para in report_text.split('\n'):
                    p = doc.add_paragraph(para)
                    if para.startswith('===') or para.startswith('-'):
                        p.runs[0].font.bold = True
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    elif para.startswith('✅') or para.startswith('⚠️'):
                        p.runs[0].font.bold = True
                # Save document
                doc.save(file_path)
                messagebox.showinfo("Sukses", f"Laporan berhasil disimpan ke:\n{file_path}")
            except Exception as e:
                messagebox.showerror("Error", f"Gagal menyimpan file DOCX:\n{str(e)}")

    def _save_to_excel(self):
        """Save report to Excel file"""
        if not self.analyzer.has_data and len(self.analyzer.all_senders) == 0:
            messagebox.showwarning("Peringatan", "Tidak ada data untuk disimpan!")
            return
        file_path = filedialog.asksaveasfilename(defaultextension=".xlsx",
                                               filetypes=[("Excel Files", "*.xlsx"), ("All Files", "*.*")],
                                               title="Simpan Laporan sebagai Excel")
        if file_path:
            try:
                # Create a new workbook
                wb = openpyxl.Workbook()
                ws = wb.active
                ws.title = "Chat Analysis"
                # Add title
                ws['A1'] = "LAPORAN ANALISIS PESAN KASAR WHATSAPP"
                ws['A1'].font = Font(bold=True, size=16)
                ws.merge_cells('A1:D1')
                # Add summary data
                ws.append(["Total pesan bermasalah", len(self.analyzer.offensive_messages)])
                ws.append(["Total kata kasar terdeteksi", sum(self.analyzer.word_stats.values())])
                ws.append(["Total peserta dalam chat", len(self.analyzer.all_senders)])
                ws.append([])
                # Add sender statistics
                ws.append(["Pengirim", "Total Pesan", "Pesan Kasar", "Persentase Kasar"])
                for sender, total in self.analyzer.all_senders.items():
                    offensive = self.analyzer.sender_stats.get(sender, 0)
                    percentage = (offensive / total * 100) if total > 0 else 0
                    ws.append([sender, total, offensive, f"{percentage:.1f}%"])
                # Format header
                for cell in ws[1]:
                    cell.font = Font(bold=True)
                # Add word statistics if available
                if self.analyzer.has_data:
                    ws2 = wb.create_sheet(title="Kata Kasar")
                    ws2.append(["Kata Kasar", "Frekuensi"])
                    for word, count in sorted(self.analyzer.word_stats.items(), key=lambda x: x[1], reverse=True):
                        ws2.append([word, count])
                    # Format header
                    for cell in ws2[1]:
                        cell.font = Font(bold=True)
                # Save file
                wb.save(file_path)
                messagebox.showinfo("Sukses", f"Laporan berhasil disimpan ke:\n{file_path}")
            except Exception as e:
                messagebox.showerror("Error", f"Gagal menyimpan file Excel:\n{str(e)}")

    def _save_to_pdf(self):
        """Save report to PDF file"""
        if not self.analyzer.has_data and len(self.analyzer.all_senders) == 0:
            messagebox.showwarning("Peringatan", "Tidak ada data untuk disimpan!")
            return
        file_path = filedialog.asksaveasfilename(defaultextension=".pdf",
                                               filetypes=[("PDF Files", "*.pdf"), ("All Files", "*.*")],
                                               title="Simpan Laporan sebagai PDF")
        if file_path:
            try:
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                report_text = self.analyzer.generate_report()
                lines = report_text.split('\n')
                for line in lines:
                    pdf.cell(200, 10, txt=line, ln=True, align='L')
                pdf.output(file_path)
                messagebox.showinfo("Sukses", f"Laporan berhasil disimpan ke:\n{file_path}")
            except Exception as e:
                messagebox.showerror("Error", f"Gagal menyimpan file PDF:\n{str(e)}")

    def _monitor_live_chat(self):
        """Monitor live chat for offensive language"""
        while True:
            try:
                # Check for new messages in the queue
                if not self.monitor_queue.empty():
                    message = self.monitor_queue.get()
                    sender, text = message['sender'], message['text']
                    
                    # Check for offensive language
                    matches = []
                    for word in self.analyzer.bad_words:
                        pattern = re.compile(word, re.IGNORECASE)
                        if pattern.search(text):
                            found_word = pattern.search(text).group(0)
                            matches.append(found_word)
                    
                    if matches:
                        # Show alert
                        self.after(0, self._show_alert, sender, text, matches)
            except Exception as e:
                print(f"Error in monitoring thread: {e}")
            time.sleep(1)  # Check every second

    def _show_alert(self, sender, message, bad_words):
        """Show alert for offensive language detection"""
        alert_window = tk.Toplevel(self)
        alert_window.title("Peringatan: Kata Kasar Terdeteksi!")
        alert_window.geometry("500x300")
        
        tk.Label(alert_window, 
                text="⚠️ KATA KASAR TERDETEKSI ⚠️", 
                font=('Arial', 14, 'bold'),
                fg='red').pack(pady=10)
        
        tk.Label(alert_window, 
                text=f"Pengirim: {sender}", 
                font=('Arial', 12)).pack(pady=5)
        
        tk.Label(alert_window, 
                text=f"Pesan: {message}", 
                font=('Arial', 10),
                wraplength=400).pack(pady=5)
        
        tk.Label(alert_window, 
                text=f"Kata terdeteksi: {', '.join(bad_words)}",
                font=('Arial', 10, 'bold'),
                fg='red').pack(pady=10)
        
        # Feedback edukatif
        tk.Label(alert_window,
                text='Gunakan bahasa yang lebih sopan untuk menjaga suasana positif.',
                font=('Arial', 10),
                fg='green').pack(pady=5)
        
        tk.Button(alert_window,
                text="Tutup",
                command=alert_window.destroy).pack(pady=10)

    def add_live_message(self, sender, message):
        """Add a live message to be monitored"""
        self.monitor_queue.put({'sender': sender, 'text': message})

    def _show_metrics_chart(self):
        """Show evaluation metrics chart"""
        if not hasattr(self.analyzer, 'all_senders') or len(self.analyzer.all_senders) == 0:
            messagebox.showinfo("Informasi", "Tidak ada data untuk ditampilkan! Silakan buka file chat terlebih dahulu.")
            return
        self.analyzer.show_metrics_chart()

    def _show_word_frequency_chart(self):
        """Show word frequency chart"""
        if not hasattr(self.analyzer, 'all_senders') or len(self.analyzer.all_senders) == 0:
            messagebox.showinfo("Informasi", "Tidak ada data untuk ditampilkan! Silakan buka file chat terlebih dahulu.")
            return
        self.analyzer.show_word_frequency_chart()

    def _show_sentiment_chart(self):
        """Show sentiment analysis chart"""
        if not hasattr(self.analyzer, 'all_senders') or len(self.analyzer.all_senders) == 0:
            messagebox.showinfo("Informasi", "Tidak ada data untuk ditampilkan! Silakan buka file chat terlebih dahulu.")
            return
        self.analyzer.show_sentiment_chart()

# Add multi-language offensive words to the WhatsAppAnalyzer class
def add_multi_language_support(analyzer):
    # Add English offensive words
    english_words = [
        r'\bfuck\b', r'\bshit\b', r'\basshole\b', r'\bbitch\b', r'\bcunt\b',
        r'\bdick\b', r'\bpussy\b', r'\bwhore\b', r'\bslut\b', r'\bbastard\b',
        r'\bmotherfucker\b', r'\bfucker\b', r'\bshit\b', r'\bdamn\b', r'\bcrap\b',
        r'\bretard\b', r'\bidiot\b', r'\bstupid\b', r'\bdumb\b', r'\bmoron\b'
    ]
    
    # Add Spanish offensive words
    spanish_words = [
        r'\bputa\b', r'\bputo\b', r'\bcoño\b', r'\bcabron\b', r'\bmaricón\b',
        r'\bmierda\b', r'\bverga\b', r'\bpendejo\b', r'\bchingar\b', r'\bpinche\b',
        r'\bgilipollas\b', r'\bjoder\b', r'\bmalparido\b', r'\bculero\b', r'\bzorra\b'
    ]
    
    # Add French offensive words
    french_words = [
        r'\bputain\b', r'\bmerde\b', r'\bsalope\b', r'\bconnard\b', r'\bcon\b',
        r'\bpute\b', r'\benculé\b', r'\bta gueule\b', r'\bniaiseux\b', r'\bcrisse\b',
        r'\bcalisse\b', r'\bosti\b', r'\btabarnak\b', r'\bsacrament\b', r'\bviarge\b'
    ]

    # Chinese (Mandarin) offensive words
    chinese_words = [
        r'\b傻逼\b', r'\b他妈的\b', r'\b操你妈\b', r'\b狗屎\b', r'\b王八蛋\b', r'\b贱人\b',
        r'\b去你妈的\b', r'\b死吧\b', r'\b滚蛋\b', r'\b婊子\b', r'\b操你\b', r'\b草泥马\b',
        r'\b妈的\b', r'\b屌你\b', r'\b你个死鬼\b', r'\b你妹\b', r'\b操你全家\b', r'\b死全家\b',
        r'\b臭不要脸\b', r'\b鸡巴\b', r'\b吊死你\b', r'\b狗娘养的\b', r'\b去死\b'
    ]
    
    # Add all to the analyzer's bad_words list
    analyzer.bad_words.extend(english_words)
    analyzer.bad_words.extend(spanish_words)
    analyzer.bad_words.extend(french_words)
    analyzer.bad_words.extend(chinese_words)
    
    # Recompile the pattern
    analyzer.pattern = re.compile('|'.join(analyzer.bad_words), flags=re.IGNORECASE)

if __name__ == "__main__":
    app = AnalyzerGUI()
    add_multi_language_support(app.analyzer)
    app.mainloop()
